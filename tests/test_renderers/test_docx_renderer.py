import io
import os
import tempfile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from renderers.docx_renderer import render_docx_bytes
from schemas import DocumentStyle

# Smallest valid 1x1 PNG -- python-docx needs a real image header to add a picture.
_PNG_BYTES = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
    "0000000a4944415478da637800010005000105a3a37a0000000049454e44ae426082"
)

DATA = {"sections": [{"type": "paragraph", "text": "hi", "indentation": 7}]}

# OOXML stores lengths in twips (1/20 pt); mm -> EMU -> twips -> EMU round-trips
# lose a fraction of a mm, so length comparisons use a tolerance rather than
# exact equality.
_TOLERANCE_MM = 0.1


def _mm(length) -> float:
    return int(length) / 36000


def _assert_mm(length, expected_mm: float) -> None:
    assert abs(_mm(length) - expected_mm) < _TOLERANCE_MM


@pytest.fixture
def png_path():
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tf:
        tf.write(_PNG_BYTES)
        name = tf.name
    yield name
    os.unlink(tf.name)


def _find_image_paragraph(doc):
    for p in doc.paragraphs:
        if "graphicData" in p._p.xml:
            return p
    return None


def test_default_margin():
    out = render_docx_bytes(DATA, "T", None, DocumentStyle())
    doc = Document(io.BytesIO(out))
    section = doc.sections[0]
    _assert_mm(section.top_margin, 12)
    _assert_mm(section.right_margin, 12)
    _assert_mm(section.bottom_margin, 16)
    _assert_mm(section.left_margin, 12)


def test_custom_margin():
    style = DocumentStyle(margin={"top": 5, "right": 6, "bottom": 7, "left": 8})
    out = render_docx_bytes(DATA, "T", None, style)
    doc = Document(io.BytesIO(out))
    section = doc.sections[0]
    _assert_mm(section.top_margin, 5)
    _assert_mm(section.right_margin, 6)
    _assert_mm(section.bottom_margin, 7)
    _assert_mm(section.left_margin, 8)


@pytest.mark.parametrize(
    "position,expected_alignment",
    [
        ("top", WD_ALIGN_PARAGRAPH.CENTER),
        ("bottom", WD_ALIGN_PARAGRAPH.CENTER),
        ("left", WD_ALIGN_PARAGRAPH.LEFT),
        ("right", WD_ALIGN_PARAGRAPH.RIGHT),
    ],
)
def test_image_alignment_matches_position(png_path, position, expected_alignment):
    style = DocumentStyle(image_position=position)
    out = render_docx_bytes(DATA, "T", png_path, style)
    doc = Document(io.BytesIO(out))
    image_paragraph = _find_image_paragraph(doc)
    assert image_paragraph is not None
    assert image_paragraph.alignment == expected_alignment


def test_bottom_position_places_image_after_sections(png_path):
    style = DocumentStyle(image_position="bottom")
    out = render_docx_bytes(DATA, "T", png_path, style)
    doc = Document(io.BytesIO(out))
    texts = [p.text for p in doc.paragraphs]
    image_index = next(i for i, p in enumerate(doc.paragraphs) if "graphicData" in p._p.xml)
    section_index = texts.index("hi")
    assert section_index < image_index


def test_document_default_indentation_applied():
    data = {"sections": [{"type": "paragraph", "text": "hi"}]}  # no per-section override
    style = DocumentStyle(indentation=9)
    out = render_docx_bytes(data, "T", None, style)
    doc = Document(io.BytesIO(out))
    p = next(p for p in doc.paragraphs if p.text == "hi")
    _assert_mm(p.paragraph_format.left_indent, 9)


def test_section_indentation_overrides_document_default():
    out = render_docx_bytes(DATA, "T", None, DocumentStyle(indentation=2))
    doc = Document(io.BytesIO(out))
    p = next(p for p in doc.paragraphs if p.text == "hi")
    _assert_mm(p.paragraph_format.left_indent, 7)


def test_no_indentation_leaves_left_indent_unset():
    data = {"sections": [{"type": "paragraph", "text": "hi"}]}
    out = render_docx_bytes(data, "T", None, DocumentStyle())
    doc = Document(io.BytesIO(out))
    p = next(p for p in doc.paragraphs if p.text == "hi")
    assert p.paragraph_format.left_indent is None


def test_default_heading_level_is_heading_2():
    data = {"sections": [{"heading": "H", "type": "paragraph", "text": "hi"}]}
    out = render_docx_bytes(data, "T", None, DocumentStyle())
    doc = Document(io.BytesIO(out))
    p = next(p for p in doc.paragraphs if p.text == "H")
    assert p.style.name == "Heading 2"


def test_custom_heading_level():
    data = {"sections": [{"heading": "H", "type": "paragraph", "text": "hi", "heading_level": 5}]}
    out = render_docx_bytes(data, "T", None, DocumentStyle())
    doc = Document(io.BytesIO(out))
    p = next(p for p in doc.paragraphs if p.text == "H")
    assert p.style.name == "Heading 5"


def test_default_list_uses_list_bullet_style():
    data = {"sections": [{"type": "list", "items": ["a"]}]}
    out = render_docx_bytes(data, "T", None, DocumentStyle())
    doc = Document(io.BytesIO(out))
    p = next(p for p in doc.paragraphs if p.text == "a")
    assert p.style.name == "List Bullet"


def test_ordered_list_uses_list_number_style():
    data = {"sections": [{"type": "list", "items": ["a"], "ordered": True}]}
    out = render_docx_bytes(data, "T", None, DocumentStyle())
    doc = Document(io.BytesIO(out))
    p = next(p for p in doc.paragraphs if p.text == "a")
    assert p.style.name == "List Number"
