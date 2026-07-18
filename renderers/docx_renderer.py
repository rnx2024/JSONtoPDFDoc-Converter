import contextlib
import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Inches, Mm

from schemas import DocumentStyle

_ALIGNMENT_BY_IMAGE_POSITION = {
    "top": WD_ALIGN_PARAGRAPH.CENTER,
    "bottom": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _add_image(doc: Any, img_bytes: bytes, image_position: str) -> None:
    # DOCX has no CSS-style float/text-wrap; left/right only align the image,
    # they don't wrap surrounding text around it like the PDF renderer does.
    with contextlib.suppress(UnrecognizedImageError, ValueError):
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
        doc.paragraphs[-1].alignment = _ALIGNMENT_BY_IMAGE_POSITION.get(
            image_position, WD_ALIGN_PARAGRAPH.CENTER
        )


def _resolve_indent(section: dict[str, Any], default_indentation: float) -> float:
    indent = section.get("indentation")
    return default_indentation if indent is None else indent


def _apply_indent(paragraph: Any, indent: float) -> None:
    if indent:
        paragraph.paragraph_format.left_indent = Mm(indent)


def render_docx_bytes(
    data: dict[str, Any],
    title: str,
    img_bytes: bytes | None,
    style: DocumentStyle,
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(style.margin.top)
    section.right_margin = Mm(style.margin.right)
    section.bottom_margin = Mm(style.margin.bottom)
    section.left_margin = Mm(style.margin.left)

    doc.add_heading(title, level=1)

    if img_bytes and style.image_position != "bottom":
        _add_image(doc, img_bytes, style.image_position)

    if isinstance(data, dict) and "sections" in data and isinstance(data["sections"], list):
        for s in data["sections"]:
            heading = s.get("heading")
            typ = s.get("type")
            if heading:
                doc.add_heading(str(heading), level=s.get("heading_level") or 2)
            if typ == "paragraph":
                p = doc.add_paragraph(str(s.get("text", "")))
                _apply_indent(p, _resolve_indent(s, style.indentation))
            elif typ == "list":
                indent = _resolve_indent(s, style.indentation)
                list_style = "List Number" if s.get("ordered") else "List Bullet"
                for it in s.get("items") or []:
                    p = doc.add_paragraph(str(it), style=list_style)
                    _apply_indent(p, indent)
            elif typ == "table":
                rows = s.get("rows") or []
                if rows:
                    cols = max(len(r) for r in rows)
                    t = doc.add_table(rows=0, cols=cols)
                    for r in rows:
                        cells = t.add_row().cells
                        for i, val in enumerate(r):
                            cells[i].text = str(val)
    else:
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells
        hdr[0].text = "Key"
        hdr[1].text = "Value"
        for k, v in (data or {}).items():
            if isinstance(v, dict):
                sub = doc.add_table(rows=1, cols=2)
                h = sub.rows[0].cells
                h[0].text = str(k)
                h[1].text = ""
                for kk, vv in v.items():
                    r = sub.add_row().cells
                    r[0].text = str(kk)
                    r[1].text = str(vv)
            else:
                r = t.add_row().cells
                r[0].text = str(k)
                r[1].text = str(v)

    if img_bytes and style.image_position == "bottom":
        _add_image(doc, img_bytes, style.image_position)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
